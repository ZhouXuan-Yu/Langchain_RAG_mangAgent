"""文档处理器 — 支持 PDF/DOCX/TXT/Markdown/CSV/图片/代码 的文本提取与分块."""

from __future__ import annotations

import io
import logging
import os
import re
import sqlite3
import uuid
import zipfile
import xml.etree.ElementTree as ET
from contextlib import contextmanager
from datetime import datetime
from pathlib import Path
from typing import BinaryIO, Optional

from src.config import CHROMA_PATH

logger = logging.getLogger(__name__)

PROJECT_ROOT = Path(__file__).parent.parent
DATA_DIR = PROJECT_ROOT / "data"
DOCS_DIR = DATA_DIR / "documents"
DOCS_DIR.mkdir(parents=True, exist_ok=True)
DOCS_DB = DATA_DIR / "documents.db"


def _uuid() -> str:
    return "doc_" + uuid.uuid4().hex[:12]


# ═══════════════════════════════════════════════════════════════════════════════
#  SQLite 连接
# ═══════════════════════════════════════════════════════════════════════════════

_docs_lock = sqlite3


@contextmanager
def _conn():
    conn = sqlite3.connect(str(DOCS_DB), check_same_thread=False)
    conn.row_factory = sqlite3.Row
    try:
        yield conn
        conn.commit()
    finally:
        conn.close()


def _init_db():
    with _conn() as c:
        c.execute("""
            CREATE TABLE IF NOT EXISTS documents (
                id           TEXT PRIMARY KEY,
                filename     TEXT NOT NULL,
                doc_type     TEXT NOT NULL,
                size_bytes   INTEGER DEFAULT 0,
                chunk_count  INTEGER DEFAULT 0,
                status       TEXT DEFAULT 'uploading',
                error        TEXT,
                storage_path TEXT,
                uploaded_at  TEXT
            )
        """)
        c.execute("""
            CREATE TABLE IF NOT EXISTS doc_chunks (
                id          TEXT PRIMARY KEY,
                doc_id      TEXT NOT NULL,
                chunk_index INTEGER DEFAULT 0,
                chunk_text  TEXT NOT NULL,
                FOREIGN KEY (doc_id) REFERENCES documents(id) ON DELETE CASCADE
            )
        """)
        c.execute("""
            CREATE INDEX IF NOT EXISTS idx_chunks_doc ON doc_chunks(doc_id)
        """)


_init_db()


# ═══════════════════════════════════════════════════════════════════════════════
#  文档类型检测
# ═══════════════════════════════════════════════════════════════════════════════

EXT_TO_TYPE = {
    ".pdf": "pdf",
    ".docx": "docx",
    ".doc": "docx",
    ".txt": "txt",
    ".md": "markdown",
    ".markdown": "markdown",
    ".csv": "csv",
    ".py": "code",
    ".js": "code",
    ".ts": "code",
    ".java": "code",
    ".cpp": "code",
    ".c": "code",
    ".go": "code",
    ".rs": "code",
    ".sh": "code",
    ".sql": "code",
    ".json": "code",
    ".yaml": "code",
    ".yml": "code",
    ".xml": "code",
    ".html": "code",
    ".css": "code",
    ".xlsx": "excel",
    ".xls": "excel",
    ".png": "image",
    ".jpg": "image",
    ".jpeg": "image",
    ".gif": "image",
    ".webp": "image",
    ".bmp": "image",
}


def detect_doc_type(filename: str) -> str:
    """根据文件扩展名检测文档类型."""
    _, ext = os.path.splitext(filename.lower())
    return EXT_TO_TYPE.get(ext, "txt")


ALLOWED_TYPES = set(EXT_TO_TYPE.values())


def is_allowed_type(doc_type: str) -> bool:
    return doc_type in ALLOWED_TYPES


# 旧版 Word 二进制（.doc）文件头；扩展名常为 .docx 但无法用 python-docx / OOXML 解析
_OLE_DOC_MAGIC = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"


def is_legacy_word_doc(content: bytes) -> bool:
    """是否为 Microsoft Word 97–2003 二进制文档（非 ZIP 结构的 .docx）."""
    return bool(content) and len(content) >= 8 and content[:8] == _OLE_DOC_MAGIC


_UTF8_BOM = b"\xef\xbb\xbf"
_ZIP_UPLOAD_DOC_TYPES = frozenset({"docx", "excel"})


def normalize_zip_upload_bytes(content: bytes, doc_type: str) -> bytes:
    """
    部分工具会在 OOXML（ZIP）文件前写入 UTF-8 BOM，导致 ZipFile 无法打开。
    仅对 docx / excel（xlsx）做剥离，避免误伤其它二进制类型。
    """
    if doc_type not in _ZIP_UPLOAD_DOC_TYPES or not content:
        return content
    if content.startswith(_UTF8_BOM):
        return content[len(_UTF8_BOM) :]
    return content


def is_docx_zip_blob(content: bytes) -> bool:
    """是否为有效 ZIP 容器（标准 .docx / .xlsx；兼容前置 UTF-8 BOM）."""
    if not content:
        return False
    body = content[len(_UTF8_BOM) :] if content.startswith(_UTF8_BOM) else content
    if len(body) < 4:
        return False
    return zipfile.is_zipfile(io.BytesIO(body))


# ═══════════════════════════════════════════════════════════════════════════════
#  文本提取
# ═══════════════════════════════════════════════════════════════════════════════

def extract_text(content: bytes, doc_type: str) -> str:
    """
    根据文档类型从二进制内容中提取纯文本。

    参数:
        content: 文件二进制内容
        doc_type: 文档类型标识符

    返回:
        提取的纯文本字符串
    """
    if not content:
        return ""

    try:
        if doc_type == "pdf":
            return _extract_pdf(content)
        elif doc_type == "docx":
            return _extract_docx(content)
        elif doc_type == "txt":
            return content.decode("utf-8", errors="replace")
        elif doc_type == "markdown":
            return content.decode("utf-8", errors="replace")
        elif doc_type == "csv":
            return _extract_csv(content)
        elif doc_type == "excel":
            return _extract_excel(content)
        elif doc_type == "image":
            return _extract_image_description(content, filename="")
        elif doc_type == "code":
            return _extract_code(content)
        else:
            return content.decode("utf-8", errors="replace")
    except Exception as e:
        logger.warning(f"[doc_processor] extract failed for type {doc_type}: {e}")
        return ""


def pdf_is_password_protected(content: bytes) -> bool:
    """检测 PDF 是否需要密码（无法解密则无法提取正文）."""
    if not content:
        return False
    try:
        import fitz

        doc = fitz.open(stream=content, filetype="pdf")
        try:
            return bool(doc.needs_pass)
        finally:
            doc.close()
    except Exception:
        pass
    try:
        import pypdf

        reader = pypdf.PdfReader(io.BytesIO(content))
        return bool(getattr(reader, "is_encrypted", False))
    except Exception:
        return False


def pdf_upload_failure_hint(content: bytes) -> str:
    """提取失败时给用户的补充说明."""
    if pdf_is_password_protected(content):
        return "该 PDF 已加密或需要密码，请先解除限制或「打印为 PDF」导出后再上传。"
    return (
        "未识别到可抽取的文字层：常见于扫描版（纯图片）PDF。"
        "请使用 Word/WPS 打开后另存为带文字层的 PDF，或使用 OCR 工具识别后再上传。"
    )


def _extract_pdf(content: bytes) -> str:
    """
    依次尝试 PyMuPDF → pypdf → pdfplumber，取最长结果。
    PyMuPDF 对中文、学术排版、嵌入字体 PDF 兼容性通常最好。
    """
    methods: list[tuple[str, str]] = []

    # 1. PyMuPDF（优先）
    try:
        import fitz

        doc = fitz.open(stream=content, filetype="pdf")
        try:
            if doc.needs_pass:
                logger.warning("[pdf] PyMuPDF: document requires password")
                return ""
            pieces: list[str] = []
            for page in doc:
                t = page.get_text()
                if t and t.strip():
                    pieces.append(t)
            joined = "\n\n".join(pieces)
            if joined.strip():
                methods.append(("pymupdf", joined))
                logger.info("[pdf] pymupdf extracted %s pages, %s chars", doc.page_count, len(joined))
            else:
                logger.debug("[pdf] pymupdf: 0 chars (may be image-only or odd encoding)")
        finally:
            doc.close()
    except ImportError:
        logger.warning("[pdf] pymupdf (fitz) not installed — run: pip install pymupdf")
    except Exception as e:
        logger.warning("[pdf] pymupdf error: %s", e, exc_info=True)

    # 2. pypdf
    try:
        import pypdf

        reader = pypdf.PdfReader(io.BytesIO(content))
        if getattr(reader, "is_encrypted", False):
            try:
                if reader.decrypt("") == 0:
                    logger.warning("[pdf] pypdf: encrypted PDF, cannot decrypt without password")
                    if not methods:
                        return ""
            except Exception:
                if not methods:
                    return ""
        pypdf_texts: list[str] = []
        for page in reader.pages:
            t = page.extract_text()
            if t:
                pypdf_texts.append(t)
        joined = "\n\n".join(pypdf_texts)
        if joined.strip():
            methods.append(("pypdf", joined))
            logger.info("[pdf] pypdf extracted %s pages, %s chars", len(pypdf_texts), len(joined))
        else:
            logger.debug("[pdf] pypdf extracted 0 chars")
    except ImportError:
        logger.warning("[pdf] pypdf not installed — run: pip install pypdf")
    except Exception as e:
        logger.warning("[pdf] pypdf error: %s", e)

    # 3. pdfplumber
    try:
        import pdfplumber

        with pdfplumber.open(io.BytesIO(content)) as pdf:
            plumb_texts: list[str] = []
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    plumb_texts.append(t)
        joined = "\n\n".join(plumb_texts)
        if joined.strip():
            methods.append(("pdfplumber", joined))
            logger.info("[pdf] pdfplumber extracted %s pages, %s chars", len(plumb_texts), len(joined))
        else:
            logger.debug("[pdf] pdfplumber extracted 0 chars")
    except ImportError:
        logger.debug("[pdf] pdfplumber not installed")
    except Exception as e:
        logger.warning("[pdf] pdfplumber error: %s", e)

    if not methods:
        logger.warning("[pdf] all extractors returned empty — likely image-only/scanned PDF or unsupported encoding")
        return ""

    best_lib, best_text = max(methods, key=lambda x: len(x[1]))
    logger.info("[pdf] using %s result (%s chars)", best_lib, len(best_text))
    return best_text


_W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _extract_docx_ooxml(content: bytes) -> str:
    """
    直接从 OOXML（ZIP 内 word/*.xml）收集所有 w:t 文本。
    可覆盖：文本框、部分 SmartArt、页眉页脚、脚注等 python-docx 未完整暴露的结构。
    """
    pieces: list[str] = []
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as zf:
            for name in sorted(zf.namelist()):
                if not name.startswith("word/") or not name.endswith(".xml"):
                    continue
                if "/media/" in name or "/embeddings/" in name or "/activeX/" in name:
                    continue
                try:
                    raw = zf.read(name)
                    root = ET.fromstring(raw)
                except (ET.ParseError, KeyError, OSError) as e:
                    logger.debug("[docx] skip xml %s: %s", name, e)
                    continue
                for el in root.iter(_W_NS + "t"):
                    if el.text:
                        pieces.append(el.text)
                    if el.tail:
                        pieces.append(el.tail)
        text = "".join(pieces)
        text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", text)
        text = re.sub(r"\n{3,}", "\n\n", text).strip()
        if text:
            logger.info("[docx] OOXML fallback collected %s chars from word/*.xml", len(text))
        return text
    except zipfile.BadZipFile:
        logger.warning("[doc_processor] DOCX is not a valid ZIP (corrupt or not .docx)")
        return ""
    except Exception as e:
        logger.warning(f"[doc_processor] DOCX OOXML extraction error: {e}", exc_info=True)
        return ""


def _extract_docx_via_python_docx(content: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(content))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    hf_lines: list[str] = []
    for section in doc.sections:
        for part in (section.header, section.footer):
            if part is None:
                continue
            for p in part.paragraphs:
                t = p.text.strip()
                if t:
                    hf_lines.append(t)
        if section.different_first_page_header_footer:
            for p in section.first_page_header.paragraphs:
                t = p.text.strip()
                if t:
                    hf_lines.append(t)
            for p in section.first_page_footer.paragraphs:
                t = p.text.strip()
                if t:
                    hf_lines.append(t)

    tables_text: list[str] = []
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                tables_text.append(" | ".join(row_text))

    parts_out: list[str] = []
    if hf_lines:
        parts_out.append("\n".join(hf_lines))
    parts_out.append("\n".join(paragraphs))
    if tables_text:
        parts_out.append("[表格内容]\n" + "\n".join(tables_text))
    result = "\n\n".join(s for s in parts_out if s.strip())
    logger.info(
        "[docx] python-docx: %s body paras, %s hf lines, %s table rows, %s chars",
        len(paragraphs),
        len(hf_lines),
        len(tables_text),
        len(result),
    )
    return result


def _merge_docx_extractions(api_text: str, ooxml_text: str) -> str:
    a = (api_text or "").strip()
    x = (ooxml_text or "").strip()
    if not a:
        return ooxml_text or ""
    if not x:
        return api_text or ""
    if x in a:
        return api_text or ""
    if a in x:
        return ooxml_text or ""
    # OOXML 明显更完整（常见于文本框 / 复杂排版论文）
    if len(x) > len(a) * 1.08:
        return ooxml_text or ""
    return (api_text or "") + "\n\n" + x


def _extract_docx(content: bytes) -> str:
    content = normalize_zip_upload_bytes(content, "docx")
    if is_legacy_word_doc(content):
        logger.warning("[doc_processor] file is legacy .doc (OLE), not Office Open XML — cannot parse as docx")
        return ""
    if not is_docx_zip_blob(content):
        logger.warning("[doc_processor] DOCX content does not start with ZIP signature PK — file may be corrupt or mislabeled")
        # 仍尝试 python-docx，少数容器可能例外
    ooxml_text = ""
    if is_docx_zip_blob(content):
        ooxml_text = _extract_docx_ooxml(content)

    try:
        api_text = _extract_docx_via_python_docx(content)
    except ImportError:
        logger.warning("[doc_processor] python-docx not installed, using OOXML only")
        return ooxml_text
    except Exception as e:
        logger.warning(f"[doc_processor] DOCX python-docx extraction error: {e}", exc_info=True)
        api_text = ""

    merged = _merge_docx_extractions(api_text, ooxml_text)
    return merged


def _extract_csv(content: bytes) -> str:
    try:
        import pandas as pd

        df = pd.read_csv(io.BytesIO(content))
        return df.to_string(index=False)
    except ImportError:
        # 降级：逐行读取
        try:
            text = content.decode("utf-8", errors="replace")
            return text[:5000]  # 截断避免过长
        except Exception:
            return ""
    except Exception as e:
        logger.warning(f"[doc_processor] CSV extraction error: {e}")
        return ""


def _extract_excel(content: bytes) -> str:
    """提取 Excel 文件的文本内容（支持 .xlsx 和 .xls）."""
    content = normalize_zip_upload_bytes(content, "excel")
    try:
        import pandas as pd

        # 尝试读取所有 sheet
        excel_file = pd.ExcelFile(io.BytesIO(content))
        sheets_text = []
        for sheet_name in excel_file.sheet_names:
            try:
                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                sheet_content = df.to_string(index=False)
                if sheet_content.strip():
                    sheets_text.append(f"[Sheet: {sheet_name}]\n{sheet_content}")
            except Exception as e:
                logger.warning(f"[excel] failed to read sheet {sheet_name}: {e}")
        result = "\n\n".join(sheets_text)
        logger.info(f"[excel] extracted {len(excel_file.sheet_names)} sheets, total {len(result)} chars")
        return result
    except ImportError:
        logger.warning("[doc_processor] pandas/openpyxl not installed for Excel extraction")
        return ""
    except Exception as e:
        logger.warning(f"[doc_processor] Excel extraction error: {e}", exc_info=True)
        return ""


def _extract_image_description(content: bytes, filename: str) -> str:
    """
    图片暂存，返回占位符文本。
    实际描述由 Agent 的 process_image 工具生成。
    """
    # 将图片保存到临时路径，由 Agent 处理
    tmp_path = DOCS_DIR / f"img_{uuid.uuid4().hex[:8]}.tmp"
    tmp_path.write_bytes(content)
    return f"[IMAGE_PLACEHOLDER: {tmp_path.name}]"


def _extract_code(content: bytes) -> str:
    """代码文件：保留原始文本，最多 5000 字符."""
    text = content.decode("utf-8", errors="replace")
    return text[:5000]


# ═══════════════════════════════════════════════════════════════════════════════
#  文本分块
# ═══════════════════════════════════════════════════════════════════════════════

CHUNK_SIZE = 500  # 每个 chunk 的字符数
CHUNK_OVERLAP = 50  # 重叠字符数


def chunk_text(text: str, chunk_size: int = CHUNK_SIZE) -> list[str]:
    """
    将长文本切分为重叠的小块。

    策略：按段落切分，单段落过长则按字符硬切。
    """
    if not text:
        return []

    # 先按段落分割
    paragraphs = text.split("\n\n")
    chunks = []
    current = ""

    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
        if len(current) + len(para) + 2 <= chunk_size:
            current += ("\n\n" if current else "") + para
        else:
            if current:
                chunks.append(current)
            # 如果单个段落就超过 chunk_size，按字符硬切
            if len(para) > chunk_size:
                for i in range(0, len(para), chunk_size - CHUNK_OVERLAP):
                    sub = para[i : i + chunk_size]
                    if sub.strip():
                        chunks.append(sub)
                current = ""
            else:
                current = para

    if current:
        chunks.append(current)

    return [c.strip() for c in chunks if c.strip()]


# ═══════════════════════════════════════════════════════════════════════════════
#  DocumentStore — SQLite 文档元数据管理
# ═══════════════════════════════════════════════════════════════════════════════

class DocumentStore:
    """文档元数据 + chunk 存储."""

    def list_docs(self, doc_type: Optional[str] = None) -> list[dict]:
        with _conn() as c:
            if doc_type:
                rows = c.execute(
                    "SELECT * FROM documents WHERE doc_type=? ORDER BY uploaded_at DESC",
                    (doc_type,),
                ).fetchall()
            else:
                rows = c.execute(
                    "SELECT * FROM documents ORDER BY uploaded_at DESC"
                ).fetchall()
        return [self._row_to_dict(r) for r in rows]

    def get(self, doc_id: str) -> Optional[dict]:
        with _conn() as c:
            row = c.execute("SELECT * FROM documents WHERE id=?", (doc_id,)).fetchone()
        if not row:
            return None
        return self._row_to_dict(row)

    def save(
        self,
        filename: str,
        doc_type: str,
        size_bytes: int,
        content: bytes,
    ) -> dict:
        """保存文档：写入文件 + 记录元数据."""
        doc_id = _uuid()
        storage_name = f"{doc_id}_{filename}"
        storage_path = DOCS_DIR / storage_name
        storage_path.write_bytes(content)

        now = datetime.now().isoformat()
        with _conn() as c:
            c.execute(
                "INSERT INTO documents (id,filename,doc_type,size_bytes,status,storage_path,uploaded_at) "
                "VALUES (?,?,?,?,?,?,?)",
                (doc_id, filename, doc_type, size_bytes, "processing", str(storage_path), now),
            )

        logger.info(f"[doc_store] saved document {doc_id}: {filename}")
        return self.get(doc_id)

    def update_status(
        self,
        doc_id: str,
        status: str,
        chunk_count: Optional[int] = None,
        error: Optional[str] = None,
    ) -> Optional[dict]:
        updates = ["status=?"]
        params: list = [status]
        if chunk_count is not None:
            updates.append("chunk_count=?")
            params.append(chunk_count)
        if error is not None:
            updates.append("error=?")
            params.append(error)
        params.append(doc_id)

        with _conn() as c:
            c.execute(f"UPDATE documents SET {','.join(updates)} WHERE id=?", params)
        return self.get(doc_id)

    def delete(self, doc_id: str) -> bool:
        """删除文档：移除文件 + 删除元数据（cascade 删除 chunks）."""
        doc = self.get(doc_id)
        if not doc:
            return False
        # 删除存储文件
        if doc.get("storage_path"):
            p = Path(doc["storage_path"])
            if p.exists():
                p.unlink()
        with _conn() as c:
            c.execute("DELETE FROM documents WHERE id=?", (doc_id,))
        logger.info(f"[doc_store] deleted document {doc_id}")
        return True

    def add_chunks(self, doc_id: str, chunks: list[str]) -> int:
        """为文档存储文本块."""
        with _conn() as c:
            for idx, chunk in enumerate(chunks):
                chunk_id = f"{doc_id}_c{idx}"
                c.execute(
                    "INSERT INTO doc_chunks (id,doc_id,chunk_index,chunk_text) VALUES (?,?,?,?)",
                    (chunk_id, doc_id, idx, chunk),
                )
        return len(chunks)

    def get_chunks(self, doc_id: str) -> list[str]:
        """获取文档的所有文本块."""
        with _conn() as c:
            rows = c.execute(
                "SELECT chunk_text FROM doc_chunks WHERE doc_id=? ORDER BY chunk_index",
                (doc_id,),
            ).fetchall()
        return [r["chunk_text"] for r in rows]

    def ingest_to_chroma(self, doc_id: str) -> int:
        """
        将文档 chunks 批量写入 ChromaDB，关联 document_id tag。
        """
        from src.memory.chroma_store import ChromaMemoryStore

        chunks = self.get_chunks(doc_id)
        if not chunks:
            # 否则 status 会永远停在 processing（extract 失败 / 空文本 / 未写入 chunk）
            self.update_status(
                doc_id,
                "failed",
                chunk_count=0,
                error="未生成文本块：文件可能为空，或缺少依赖（如 Word 需 python-docx、PDF 需 pypdf）",
            )
            return 0

        store = ChromaMemoryStore()
        doc = self.get(doc_id)
        filename = doc["filename"] if doc else doc_id

        count = 0
        for chunk in chunks:
            store.upsert(
                content=chunk[:1000],  # ChromaDB 有长度限制
                metadata={
                    "category": "document",
                    "importance": 5,
                    "document_id": doc_id,
                    "filename": filename,
                    "source": "document_ingestion",
                },
            )
            count += 1

        self.update_status(doc_id, "ready", chunk_count=count)
        logger.info(f"[doc_store] ingested {count} chunks to ChromaDB for doc {doc_id}")
        return count

    def reset(self) -> None:
        """清空文档元数据和 chunks 表（危险操作）."""
        with _conn() as c:
            c.execute("DELETE FROM documents")
            c.execute("DELETE FROM doc_chunks")
            c.commit()
        # 同时重置单例
        global _doc_store
        _doc_store = None
        logger.info("[doc_store] reset complete")

    @staticmethod
    def _row_to_dict(row: sqlite3.Row) -> dict:
        return {
            "id": row["id"],
            "filename": row["filename"],
            "doc_type": row["doc_type"],
            "size_bytes": row["size_bytes"],
            "chunk_count": row["chunk_count"],
            "status": row["status"],
            "error": row["error"],
            "storage_path": row["storage_path"],
            "uploaded_at": row["uploaded_at"],
        }


# ── 全局单例 ──────────────────────────────────────────────────────────────
_doc_store: Optional[DocumentStore] = None


def get_document_store() -> DocumentStore:
    global _doc_store
    if _doc_store is None:
        _doc_store = DocumentStore()
    return _doc_store
